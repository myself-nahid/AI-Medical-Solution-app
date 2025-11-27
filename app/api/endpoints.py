from fastapi import APIRouter, UploadFile, File, Form, HTTPException
from typing import List, Optional, Dict
import json
from pydantic import parse_obj_as
import asyncio
import io
from docx import Document
import magic

from app.api.models import GenerationRequestData, GenerationWithTokenResponse, DocumentRequest
from app.services import processing_service, generation_service, token_service
from app.prompts import SectionName

router = APIRouter()


# In endpoints.py - Replace the ENTIRE process_single_file function

async def process_single_file(file: UploadFile) -> str:
    """Helper function to process one file by routing it to the correct service."""
    if not file or not file.filename:
        return ""
    
    # Read file bytes ONCE and pass them around
    file_bytes = await file.read()
    print(f"Read {len(file_bytes)} bytes from file {file.filename}")
    # Reset file pointer for any subsequent reads (defensive programming)
    await file.seek(0)
    # --- START: DETAILED LOGGING ---
    # Use python-magic to detect the true file type from its content
    detected_mime_type = magic.from_buffer(file_bytes, mime=True)
    
    print("\n--- FILE RECEIVED ---")
    print(f"  - Original Filename: {file.filename}")
    print(f"  - Client-Sent Content-Type: {file.content_type}")
    print(f"  - Server-Detected MIME Type: {detected_mime_type}")
    print("---------------------\n")
    # --- END: DETAILED LOGGING ---
    # Detect MIME type from bytes
    try:
        mime_type = magic.from_buffer(file_bytes, mime=True)
    except Exception as e:
        print(f"MIME detection failed for {file.filename}: {e}")
        mime_type = "application/octet-stream"
    
    filename_lower = file.filename.lower()
    
    # Define common file extensions
    AUDIO_EXTENSIONS = (
        '.mp3', '.wav', '.m4a', '.aac', '.flac', '.ogg', 
        '.opus', '.wma', '.aif', '.aiff', '.webm'
    )
    
    IMAGE_EXTENSIONS = (
        '.heic', '.heif', '.jpg', '.jpeg', '.png', 
        '.gif', '.bmp', '.webp', '.tiff', '.tif'
    )
    
    print(f"Processing file: {file.filename}, MIME: {mime_type}, Size: {len(file_bytes)} bytes")
    
    text = ""
    
    # Audio files - check BOTH MIME type AND extension
    if "audio" in mime_type or filename_lower.endswith(AUDIO_EXTENSIONS):
        # Special case: M4A files often detected as video/mp4
        if filename_lower.endswith('.m4a') or mime_type in ['audio/mp4', 'audio/x-m4a']:
            text = await processing_service.process_audio_with_api(file, file_bytes)
        elif "audio" in mime_type or filename_lower.endswith(AUDIO_EXTENSIONS):
            text = await processing_service.process_audio_with_api(file, file_bytes)
        else:
            text = f"[Unsupported audio format: {file.filename}]"
    
    # PDF files
    elif "pdf" in mime_type or filename_lower.endswith('.pdf'):
        text = processing_service.process_pdf_locally(file_bytes)
    
    # Image files - check BOTH MIME type AND extension
    elif "image" in mime_type or filename_lower.endswith(IMAGE_EXTENSIONS):
        text = await processing_service.process_image_with_api(file, file_bytes)
    
    # Video files with audio (some users might upload these)
    elif "video" in mime_type and filename_lower.endswith(('.mp4', '.mov', '.avi', '.mkv', '.m4a')):
        # M4A is technically a video container
        if filename_lower.endswith('.m4a'):
            text = await processing_service.process_audio_with_api(file, file_bytes)
        else:
            text = f"[Video files not supported. Please extract audio first: {file.filename}]"
    
    else:
        text = f"[Unsupported file type: {mime_type} - {file.filename}]"
    
    return f"--- START OF FILE: {file.filename} ---\n{text}\n--- END OF FILE ---\n"


async def process_files(files: List[UploadFile]) -> str:
    """Processes a list of uploaded files in parallel."""
    if not files:
        return ""
    
    # Filter out empty files
    valid_files = [f for f in files if f and f.filename]
    
    if not valid_files:
        return ""
    
    tasks = [process_single_file(file) for file in valid_files]
    results = await asyncio.gather(*tasks, return_exceptions=True)
    
    # Handle any exceptions in processing
    processed_results = []
    for i, result in enumerate(results):
        if isinstance(result, Exception):
            print(f"Error processing file {valid_files[i].filename}: {result}")
            processed_results.append(f"[Error processing {valid_files[i].filename}]")
        elif result:
            processed_results.append(result)
    
    return "\n".join(processed_results)


# --- API Endpoints (Updated with Parallel Logic) ---

@router.post("/generate_section/{section_name}", response_model=GenerationWithTokenResponse)
async def generate_section_endpoint(
    section_name: SectionName,
    files: List[UploadFile] = File(...),
    request_data: str = Form(...),
    language: str = Form(...),
    specialty: str = Form(...),
    user_id: str = Form(...)
):
    if section_name == SectionName.ANALYSIS_AND_PLAN:
        raise HTTPException(status_code=400, detail="Use the dedicated /generate_analysis_plan endpoint.")

    try:
        request_body = parse_obj_as(GenerationRequestData, json.loads(request_data))
    except Exception:
        raise HTTPException(status_code=400, detail="Invalid JSON format for request_data.")
        
    has_tokens = await token_service.check_user_tokens(user_id=user_id)
    if not has_tokens:
        raise HTTPException(status_code=402, detail="You have no tokens remaining.")

    raw_input_text = await process_files(files)
    
    # --- START: PARALLEL EXECUTION ---
    # Create two tasks that can run concurrently
    generation_task = generation_service.generate_structured_text(
        previous_sections=request_body.previous_sections,
        raw_input=raw_input_text,
        section_name=section_name.value,
        physician_notes=request_body.physician_notes, 
        language=language,
        specialty=specialty
    )
    token_reporting_task = token_service.report_and_get_remaining_tokens(user_id=user_id, amount=5)

    # Run both tasks at the same time and wait for both to complete
    results = await asyncio.gather(generation_task, token_reporting_task)
    
    generated_text = results[0]
    # remaining_token = results[1]
    # --- END: PARALLEL EXECUTION ---

    # ✅ --- START OF TOKEN LOGIC FIX ---
    remaining_token = -1
    # Only deduct tokens if the entire process was successful.
    # We check if the result does NOT contain our known error indicators.
    if not ("[Error" in raw_input_text or "[AI_ERROR]" in generated_text or "[Unsupported" in raw_input_text):
        remaining_token = await token_service.report_and_get_remaining_tokens(user_id=user_id, amount=5)
    else:
        print("✗ File processing or AI generation failed. Skipping token deduction.")
    # ✅ --- END OF TOKEN LOGIC FIX ---
    
    return GenerationWithTokenResponse(
        section_name=section_name.value, 
        generated_text=generated_text,
        remaining_token=remaining_token
    )

@router.post("/generate_analysis_plan", response_model=GenerationWithTokenResponse)
async def generate_analysis_plan_endpoint(
    files: List[UploadFile] = File(...),
    request_data: str = Form(...),
    language: str = Form(...),
    specialty: str = Form(...),
    user_id: str = Form(...)
):
    try:
        request_body = parse_obj_as(GenerationRequestData, json.loads(request_data))
    except Exception:
        raise HTTPException(status_code=400, detail="Invalid JSON format for request_data.")
        
    has_tokens = await token_service.check_user_tokens(user_id=user_id)
    if not has_tokens:
        raise HTTPException(status_code=402, detail="You have no tokens remaining.")
        
    analysis_plan_text = await process_files(files)

    # --- START: PARALLEL EXECUTION ---
    generation_task = generation_service.generate_analysis_and_plan(
        previous_sections=request_body.previous_sections,
        analysis_plan_text=analysis_plan_text,
        physician_notes=request_body.physician_notes,
        language=language,
        specialty=specialty
    )
    token_reporting_task = token_service.report_and_get_remaining_tokens(user_id=user_id, amount=5)

    results = await asyncio.gather(generation_task, token_reporting_task)
    generated_text = results[0]
    # remaining_token = results[1]
    # --- END: PARALLEL EXECUTION ---

    remaining_token = -1
    if not ("[Error" in analysis_plan_text or "[AI_ERROR]" in generated_text or "[Unsupported" in analysis_plan_text):
        remaining_token = await token_service.report_and_get_remaining_tokens(user_id=user_id, amount=5)
    else:
        print("✗ File processing or AI generation failed. Skipping token deduction.")

    return GenerationWithTokenResponse(
        section_name=SectionName.ANALYSIS_AND_PLAN.value,
        generated_text=generated_text,
        remaining_token=remaining_token
    )

@router.post("/quick_report", response_model=GenerationWithTokenResponse)
async def quick_report_endpoint(
    files: List[UploadFile] = File(...),
    language: str = Form(...),
    specialty: str = Form(...),
    user_id: str = Form(...)
):
    has_tokens = await token_service.check_user_tokens(user_id=user_id)
    if not has_tokens:
        raise HTTPException(
            status_code=402,
            detail="You have no tokens remaining"
        )
    
    extracted_text = await process_files(files)
    
    # --- START: PARALLEL EXECUTION ---
    generation_task = generation_service.generate_structured_text(
        section_name=SectionName.QUICK_REPORT.value,
        raw_input=extracted_text,
        previous_sections={},
        physician_notes="", 
        language=language,
        specialty=specialty
    )
    token_reporting_task = token_service.report_and_get_remaining_tokens(user_id=user_id, amount=5)

    results = await asyncio.gather(generation_task, token_reporting_task)
    generated_text = results[0]
    # remaining_token = results[1]
    # --- END: PARALLEL EXECUTION ---
    
    remaining_token = -1
    if not ("[Error" in extracted_text or "[AI_ERROR]" in generated_text or "[Unsupported" in extracted_text):
        remaining_token = await token_service.report_and_get_remaining_tokens(user_id=user_id, amount=5)
    else:
        print("✗ File processing or AI generation failed. Skipping token deduction.")

    return GenerationWithTokenResponse(
        section_name=SectionName.QUICK_REPORT.value,
        generated_text=generated_text,
        remaining_token=remaining_token
    )

# @router.post("/generate_document")
# async def generate_docx_endpoint(request: DocumentRequest):
#     """
#     Takes all generated section texts and creates a .docx file for download.
#     This does not cost tokens and is not parallelized.
#     """
#     document = Document()
#     document.add_heading('Clinical Note', level=1)
    
#     section_order = [
#         SectionName.PRESENT_ILLNESS, SectionName.PAST_MEDICAL_HISTORY,
#         SectionName.PHYSICAL_EXAM, SectionName.LABS_AND_IMAGING,
#         SectionName.PROPOSED_DIAGNOSIS, SectionName.ANALYSIS_AND_PLAN
#     ]
    
#     for section_enum in section_order:
#         if section_enum in request.sections:
#             document.add_heading(section_enum.value, level=2)
#             document.add_paragraph(request.sections[section_enum])
#             document.add_paragraph()
            
#     file_stream = io.BytesIO()
#     document.save(file_stream)
#     file_stream.seek(0)
    
#     return StreamingResponse(
#         file_stream, 
#         media_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
#         headers={'Content-Disposition': 'attachment; filename=clinical_note.docx'}
#     )